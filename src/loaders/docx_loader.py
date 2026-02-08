"""DOCX loader — текст + структура через python-docx."""
from pathlib import Path

from src.loaders.base import BaseLoader, normalize_text


def _extract_table(table) -> str:
    """Текстовое представление таблицы."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(" | ".join(cells))
    return "\n".join(rows)


class DOCXLoader(BaseLoader):
    extensions = (".docx",)
    mime_types = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    )

    def load(self, path: Path) -> str:
        try:
            from docx import Document
            doc = Document(str(path))
            parts = []
            for p in doc.paragraphs:
                if p.text:
                    parts.append(p.text)
            for t in doc.tables:
                parts.append(_extract_table(t))
            return normalize_text("\n".join(parts))
        except Exception:
            return ""
